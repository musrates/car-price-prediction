from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import pandas as pd

# Paths for images
IMAGES = [
    ("correlation_matrix.png", "Korelasyon Matrisi"),
    ("age_vs_price.png", "Araba Yaşı vs Fiyat (Scatter)"),
    ("model_comparison_improved.png", "Model Karşılaştırma Grafiği"),
    ("feature_importance_improved.png", "Özellik Önem Dereceleri"),
    ("error_analysis.png", "Hata Analizi Grafikleri"),
]

# Create document
doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def para(text):
    doc.add_paragraph(text)

def add_code_block(code: str):
    p = doc.add_paragraph()
    for line in code.strip("\n").split("\n"):
        r = p.add_run(line + "\n")
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)

add_title('ARABA FİYAT TAHMİN MODELİ — SINAV RAPORU')

meta = doc.add_paragraph()
meta.add_run('Tarih: ').bold = True
meta.add_run(datetime.now().strftime('%d %B %Y'))
meta.add_run('    |    En İyi Model: ').bold = True
meta.add_run('Gradient Boosting Regressor (Tuned)')
meta.add_run('    |    Final R²: ').bold = True
meta.add_run('0.7912')

# 1. Giriş
h2('1. Giriş ve Problem Tanımı')
para('Bu proje, ikinci el araçların teknik ve kullanım özelliklerinden yola çıkarak satış fiyatını tahmin etmeyi amaçlar. Problem denetimli regresyon problemidir ve hedef değişken Hint Rupisi (₹) cinsinden fiyatıdır. İş hedefi, yüksek genelleme başarımı (R² ≈ 0.79) ile güvenilir tahminler üretebilen bir model geliştirmektir.')

# 2. Veri Seti
h2('2. Veri Seti ve Özeti')
para('Ham veri 5.512 satır ve 9 temel kolondan oluşur. Temizleme ve özellik mühendisliğinden sonra 4.921 satır ve 33 özellik ile modelleme yapılmıştır. Veri seti; araç adı, fiyat (Lakh/Crore formatında metin), km, yakıt tipi, şanzıman, sahiplik, üretim yılı, motor hacmi ve koltuk sayısını içerir.')

bullet('Toplam gözlem: 5.512 → 4.921 (outlier temizliği sonrası)')
bullet('Nihai özellik sayısı: 33 (sayısal + one-hot kodlu kategorik)')
bullet('Hedef değişken: fiyat (₹)')

# Örnek veri tablosu (2 satır)
try:
    _df = pd.read_csv('car_price.csv')
    cols = ['car_name','car_prices_in_rupee','kms_driven','fuel_type','transmission','manufacture','engine','Seats']
    sample = _df[cols].head(2)
    doc.add_paragraph('Örnek Veri (2 satır):')
    t = doc.add_table(rows=1, cols=len(cols))
    for i,c in enumerate(cols):
        t.rows[0].cells[i].text = c
    for _, row in sample.iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = str(row[c])
except Exception:
    para('Örnek veri tablo oluşturulamadı (CSV okunamadı).')

# 3. Ön İşleme
h2('3. Veri Temizleme ve Dönüştürme')
para('Fiyat sütunundaki "Lakh/Crore" formatları sayısala çevrilmiş, virgül ve boşluklar temizlenmiştir. Aykırı değerler IQR yöntemi (Q1-3*IQR, Q3+3*IQR) ile temizlenmiş ve 566 satır veri çıkarılmıştır. Kategorik değişkenler (yakıt, şanzıman, marka) one-hot encoding ile temsil edilmiştir. Sahiplik sıralı olarak kodlanmıştır.')

h3('Önemli Dönüşümler')
bullet('Fiyat dönüştürme: "10.03 Lakh" → 1,003,000 ₹; "1.10 Crore" → 11,000,000 ₹')
bullet('Aykırı değer temizliği: IQR ile 566 satır çıkarıldı (genelleme arttı)')
bullet('Ölçekleme: StandardScaler (ağaç dışı modeller için)')

# 4. Özellik Mühendisliği
h2('4. Özellik Mühendisliği (Feature Engineering)')
para('Ham değişkenlerden türetilen beş anahtar özellik:')
bullet('Marka (brand): Araç adından çıkarıldı; en yüksek önem puanına sahip.')
bullet('Araba yaşı (car_age): 2025 - üretim yılı; değer kaybını yansıtır.')
bullet('Yıllık km (km_per_year): Kullanım yoğunluğunu normalize eder.')
bullet('Motor/Koltuk oranı (engine_per_seat): Segment/performans göstergesi.')
bullet('Yüksek performans bayrağı (high_performance): >2000 cc için 1.')

# 5. EDA ve Korelasyon
h2('5. Keşifsel Analiz ve Korelasyon')
para('Sayısal değişkenler arası ilişkiler incelenmiş, fiyat ile anlamlı ilişkiler tespit edilmiştir: motor hacmi pozitif, yaş/ km değişkenleri negatif korelasyon göstermektedir. Bu bulgular, doğrusal olmayan modellerin daha uygun olacağına işaret etmektedir.')

bullet('engine_numeric ↗ fiyat: +0.68 (güçlü pozitif) — daha büyük motorlar daha pahalıdır.')
bullet('engine_per_seat ↗ fiyat: +0.57 (pozitif) — segment/performans etkisi.')
bullet('car_age ↗ fiyat: −0.52 (orta negatif) — yaşlandıkça değer kaybı.')
bullet('kms_numeric ↗ fiyat: −0.39 (negatif) — kullanım arttıkça aşınma etkisi.')
bullet('km_per_year ↗ fiyat: −0.31 (negatif) — yoğun yıllık kullanım değer düşürür.')
bullet('ownership_encoded ↗ fiyat: −0.45 (negatif) — çoklu sahiplik algı değeri düşürür.')
para('Not: Korelasyon nedensellik değildir. Özellikler arası etkileşimler (ör. marka×yaş) nedeniyle doğrusal olmayan yöntemler tercih edilmiştir. Çoklu doğrusal bağlanım riski ağaç tabanlı modellerde daha iyi yönetilir.')

# Görseller
for path, caption in IMAGES:
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6))
        cap = doc.add_paragraph(f'Sekil — {caption}')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 6. Modelleme Stratejisi
h2('6. Modelleme Stratejisi ve Kıyaslama')
para('Doğrusal ve ağaç tabanlı bir dizi model denenmiştir: Linear/Ridge/Lasso, Decision Tree, Random Forest ve Gradient Boosting. 5-kat çapraz doğrulama ile model seçimi yapılmış, en iyi genel performansı Gradient Boosting göstermiştir.')

# Tablo: model karşılaştırma
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Test R²'
hdr_cells[2].text = 'RMSE (₹)'
hdr_cells[3].text = 'MAE (₹)'

rows = [
    ('Gradient Boosting', '0.7912', '279,173', '183,392'),
    ('Random Forest', '0.7531', '303,599', '193,262'),
    ('Ridge/Lasso/Linear', '≈0.709', '329,578', '227,426'),
    ('Decision Tree', '0.6579', '357,354', '219,217'),
]
for r in rows:
    cells = table.add_row().cells
    cells[0].text, cells[1].text, cells[2].text, cells[3].text = r

# 7. Neden Gradient Boosting?
h2('7. Neden Gradient Boosting En İyi? — Ayrıntılı Gerekçe')
para('Gradient Boosting (GB), zayıf öğrenicilerden (sığ karar ağaçları) oluşan bir topluluk yöntemidir. Her yeni ağaç, önceki ağaçların hatalarını hedefleyerek art arda iyileştirme yapar. Bu mekanizma, doğrusal olmayan ve etkileşim içeren ilişkileri güçlü şekilde yakalar.')

h3('Veri Özelliklerine Uyum')
bullet('Doğrusal olmayan yapı: Marka × Yaş × KM etkileşimleri doğrusaldan sapar.')
bullet('Karma değişken türleri: Sayısal + kategorik (one-hot) karışımı GB için uygundur.')
bullet('Sınırlı örnek sayısı (4.9K): GB, karmaşıklığı kademeli artırarak overfitting’i sınırlar.')

h3('Düzenleme (Regularization) ve Genel Başarım')
bullet('learning_rate=0.1: Her adımı küçük tutarak aşırı uyum riskini azaltır.')
bullet('max_depth=5: Ağaçların ifade gücü yeterli, ama aşırı ayrıntıya kaçmıyor.')
bullet('min_samples_split=2: Bölünme esnekliğini korurken CV sonuçlarıyla dengeli.')

h3('Diğer Modellerle Teknik Karşılaştırma')
bullet('Linear/Ridge/Lasso: Doğrusal varsayım → karmaşık etkileşimleri ıskalar (R² ≈ 0.71).')
bullet('Decision Tree: Tek ağaç yüksek varyans; genellemeyi zayıflatır (R² ≈ 0.66).')
bullet('Random Forest: Varyansı düşürür ama bias daha yüksek kalır; GB kadar isabetli değil (R² ≈ 0.75).')

# 8. Hiperparametre Araması
h2('8. Hiperparametre Optimizasyonu (GridSearchCV)')
para('n_estimators ∈ {100,150,200}, max_depth ∈ {5,7,10}, learning_rate ∈ {0.05,0.1,0.15}, min_samples_split ∈ {2,5} olmak üzere 54 aday yapı 5-kat CV ile değerlendirilmiştir. En iyi yapı: n_estimators=200, max_depth=5, learning_rate=0.1, min_samples_split=2. CV R²=0.7691 ± 0.026 ile test R²=0.7912 tutarlıdır.')

# 8.1 Önemli Kod Parçaları
h3('Önemli Kod: Veri Yükleme ve Fiyat Dönüştürme')
add_code_block("""
import pandas as pd

df = pd.read_csv('car_price.csv')

def convert_price(s):
    s = str(s).replace(',', '').strip()
    if 'Crore' in s: return float(s.replace('Crore','').strip())*10000000
    if 'Lakh' in s:  return float(s.replace('Lakh','').strip())*100000
    return pd.to_numeric(s, errors='coerce')

df['price_numeric'] = df['car_prices_in_rupee'].apply(convert_price)
""")

h3('Önemli Kod: Korelasyon ve Isı Haritası')
add_code_block("""
import seaborn as sns, matplotlib.pyplot as plt
num_cols = ['engine_numeric','car_age','km_per_year','kms_numeric','price_numeric']
corr = df[num_cols].corr()
sns.heatmap(corr, annot=True, fmt='.2f', cmap='coolwarm')
plt.savefig('correlation_matrix.png', dpi=300)
""")

# 9. Metrikler
h2('9. Değerlendirme Metrikleri ve Yorum')
para('• R² (Açıklama Gücü): 0 ile 1 arası; 0.79 → değişkenliğin %79’u açıklanıyor.\n• RMSE: Ortalama karekök hata; ₹279,173 → büyük hatalara duyarlı.\n• MAE: Ortalama mutlak hata; ₹183,392 → tipik sapmayı yansıtır.\n• Yüzde Hata: ≈%24.83 → fiyat büyüklüğüne göre normalleştirilmiş hata.')

# 10. Hata Analizi
h2('10. Hata Analizi ve Genelleme')
para('Artık dağılımı simetrik ve sıfır etrafında yoğunlaşmıştır. Hata yüzdelerinin medyanı düşüktür; büyük değerli lüks segmentlerde göreli hata nispeten artabilir (pazar volatilitesi). Train ve test R² birbirine yakın olduğundan aşırı uyum gözlenmemiştir. CV puanları tutarlıdır.')

# 11. Sınırlamalar ve Gelecek İşler
h2('11. Sınırlamalar ve Gelecek Çalışmalar')
bullet('Özellik kapsamı: Donanım paketleri, servis kayıtları, hasar bilgisi yok.')
bullet('Zaman etkisi: Dönemsel/ekonomik dalgalanmalar modele açıkça kodlanmadı.')
bullet('Coğrafya: Bölgesel fiyat farklılıkları veri içinde etiketlenmedi.')
bullet('Gelecek: CatBoost/XGBoost, SHAP ile açıklanabilirlik, zaman ve bölge etiketleri.')

# 12. Sonuç
h2('12. Sonuç')
para('Gradient Boosting, doğrusal olmayan ilişkilere sahip bu problemde en iyi dengesizliği yakalayarak test R²=0.7912 ve MAE≈₹183K ile en yüksek başarıyı sağlamıştır. Model üretime alınabilir seviyededir ve sağlamlaştırmak için daha zengin özellikler eklenebilir.')

# Ekler
h2('Ekler')
para('Örnek Tahminler: Maruti Swift ≈ ₹593,759; Hyundai Creta ≈ ₹1,570,017; Honda City ≈ ₹614,180.')
para('Üretilen Dosyalar: best_model.pkl, scaler.pkl, feature_names.pkl, correlation_matrix.png, model_comparison_improved.png, error_analysis.png, feature_importance_improved.png, age_vs_price.png, model_comparison_results.csv, detailed_report.txt, model.py')

# Save
output = 'PROJE_RAPORU.docx'
try:
    doc.save(output)
    print(f'✓ Word raporu oluşturuldu: {output}')
except PermissionError:
    alt_output = 'PROJE_RAPORU_v2.docx'
    doc.save(alt_output)
    print(f'⚠ PROJE_RAPORU.docx açıktı, {alt_output} olarak kaydedildi.')
